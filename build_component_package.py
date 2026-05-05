import os
import shutil
import sqlite3
from datetime import datetime
from pathlib import Path
from urllib.request import urlopen

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


PROJECT_ROOT = Path(__file__).resolve().parent
COMPONENT_ROOT = Path(r"C:\Users\Dipesh\Desktop\By Name New\Bibha\component\Executing Folder")
DATABASE_ROOT = COMPONENT_ROOT / "Implementation" / "Source Code" / "Database"
MONITORING_ROOT = COMPONENT_ROOT / "Monitoring and Controlling"
PORTFOLIO_ROOT = COMPONENT_ROOT / "Portfolio"
EVIDENCE_ROOT = COMPONENT_ROOT / "_Evidence Screenshots"


def ensure_dirs():
    for path in [
        COMPONENT_ROOT,
        DATABASE_ROOT,
        COMPONENT_ROOT / "Implementation" / "Source Code" / "Python Flask Framework",
        MONITORING_ROOT / "User Guidelines",
        MONITORING_ROOT / "Installation Guidelines",
        MONITORING_ROOT / "Testing and Evaluation",
        PORTFOLIO_ROOT / "Final Presentation Slides",
        EVIDENCE_ROOT,
    ]:
        path.mkdir(parents=True, exist_ok=True)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def save_schema():
    db_path = PROJECT_ROOT / "backend" / "instance" / "attendance.db"
    if db_path.exists():
        shutil.copy2(db_path, DATABASE_ROOT / "attendance.db")
        conn = sqlite3.connect(db_path)
        rows = conn.execute(
            "select name, sql from sqlite_master where type='table' and name not like 'sqlite_%' order by name"
        ).fetchall()
        schema = "\n\n".join(sql + ";" for _, sql in rows if sql)
        (DATABASE_ROOT / "schema.sql").write_text(schema, encoding="utf-8")
        summary = "\n".join(f"- {name}" for name, _ in rows)
        (DATABASE_ROOT / "database_summary.txt").write_text(
            "Secure Attendance System database tables\n\n" + summary + "\n",
            encoding="utf-8",
        )
    shutil.copy2(PROJECT_ROOT / "backend" / "models.py", DATABASE_ROOT / "models.py")
    shutil.copy2(PROJECT_ROOT / "backend" / "seed_demo_data.py", DATABASE_ROOT / "seed_demo_data.py")


def make_text_image(lines, output_path, title):
    width, height = 1500, 900
    image = Image.new("RGB", (width, height), "#111827")
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype("arial.ttf", 34)
        font = ImageFont.truetype("consola.ttf", 24)
    except Exception:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()
    draw.text((40, 35), title, fill="#f9fafb", font=font_title)
    y = 100
    for line in lines:
        draw.text((40, y), line, fill="#d1d5db", font=font)
        y += 34
        if y > height - 50:
            break
    image.save(output_path)


def collect_test_evidence():
    lines = [
        f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "Project: Secure ESP32 Attendance System",
        "",
    ]

    try:
        db_path = PROJECT_ROOT / "backend" / "instance" / "attendance.db"
        conn = sqlite3.connect(db_path)
        tables = conn.execute(
            "select name from sqlite_master where type='table' and name not like 'sqlite_%' order by name"
        ).fetchall()
        lines.append(f"Database file found: {db_path}")
        lines.append(f"Database table count: {len(tables)}")
        for table in tables:
            count = conn.execute(f"select count(*) from {table[0]}").fetchone()[0]
            lines.append(f"  {table[0]} rows: {count}")
    except Exception as exc:
        lines.append(f"Database verification failed: {exc}")

    try:
        import face_recognition  # noqa: F401

        lines.append("face_recognition import: PASS")
    except Exception as exc:
        lines.append(f"face_recognition import: FAIL ({exc})")

    try:
        import cv2  # noqa: F401

        lines.append("OpenCV import: PASS")
    except Exception as exc:
        lines.append(f"OpenCV import: FAIL ({exc})")

    try:
        with urlopen("http://127.0.0.1:5000/login", timeout=5) as response:
            lines.append(f"Backend /login HTTP status: {response.status}")
    except Exception as exc:
        lines.append(f"Backend /login HTTP check: not running or blocked ({exc})")

    evidence_txt = EVIDENCE_ROOT / "backend_test_evidence.txt"
    evidence_png = EVIDENCE_ROOT / "backend_test_evidence.png"
    evidence_txt.write_text("\n".join(lines) + "\n", encoding="utf-8")
    make_text_image(lines, evidence_png, "Backend and Database Test Evidence")
    return lines, evidence_png


def requirements_catalog():
    doc = Document()
    add_heading(doc, "Requirements Catalog", 0)
    doc.add_paragraph("Project: Secure ESP32-CAM Attendance System")
    doc.add_paragraph("Prepared for component submission.")

    add_heading(doc, "Functional Requirements", 1)
    add_bullets(
        doc,
        [
            "Admin can authenticate and access the attendance dashboard.",
            "Admin can create, edit, activate, and deactivate student records.",
            "Admin can upload student face images for enrollment.",
            "System stores face encodings for enrolled students.",
            "Admin can create and activate attendance sessions.",
            "System generates QR codes for active sessions.",
            "ESP32-CAM device sends heartbeat status to the backend.",
            "ESP32-CAM captures QR frames and submits them for backend verification.",
            "ESP32-CAM captures face frames and submits them for backend face matching.",
            "Backend records attendance and verification outcomes.",
            "Device page displays last status, IP address, firmware version, RSSI, and last seen time.",
            "Reports page provides attendance viewing and CSV export.",
        ],
    )

    add_heading(doc, "Non-Functional Requirements", 1)
    add_bullets(
        doc,
        [
            "The system should run on a local network using a Flask backend.",
            "Device API calls must include device ID and API key headers.",
            "Face recognition should reject unclear or multi-face enrollment images.",
            "The UI should be usable by an administrator with minimal technical steps.",
            "The database should persist students, sessions, devices, face profiles, fingerprint profiles, and attendance records.",
            "The project should be demonstrable using a laptop backend, ESP32-CAM, and supported browser.",
        ],
    )

    add_heading(doc, "Hardware Requirements", 1)
    add_bullets(
        doc,
        [
            "ESP32-CAM for camera, QR, and face capture operations.",
            "ESP32 Dev Board for separated fingerprint module operation.",
            "R307S fingerprint sensor for fingerprint enrollment and verification.",
            "USB-TTL converter for ESP32-CAM flashing and serial debugging.",
            "Stable 5V power supply for ESP32-CAM.",
            "Laptop hotspot or router-based Wi-Fi network.",
        ],
    )

    add_heading(doc, "Software Requirements", 1)
    add_bullets(
        doc,
        [
            "Python 3.11 virtual environment for backend compatibility.",
            "Flask and Flask-SQLAlchemy for web application and ORM.",
            "SQLite database for local prototype persistence.",
            "OpenCV, Pillow, qrcode, face-recognition, and dlib/dlib-bin for QR and face processing.",
            "MicroPython or Arduino firmware workflow for ESP32-CAM deployment.",
        ],
    )
    doc.save(COMPONENT_ROOT / "Requirements Catalog.docx")


def user_guidelines():
    doc = Document()
    add_heading(doc, "User Guidelines", 0)
    add_heading(doc, "Administrator Perspective", 1)
    add_bullets(
        doc,
        [
            "Start the Flask backend and open the website in a browser.",
            "Login as an administrator.",
            "Open Students and create student profiles with student code, name, email, and class.",
            "Upload one clear face image per student from the student edit page.",
            "Open Sessions and create a class session with start/end time.",
            "Activate the session and show the generated QR code to the ESP32-CAM.",
            "Review attendance records from the Attendance page and export CSV when needed.",
            "Open Devices to confirm ESP32 device heartbeat, IP address, and status.",
        ],
    )
    add_heading(doc, "End User / Student Perspective", 1)
    add_bullets(
        doc,
        [
            "Stand in front of the ESP32-CAM when attendance is open.",
            "Present the active session QR code to the ESP32-CAM.",
            "Look at the camera for face verification.",
            "Use the fingerprint device when requested by the final two-device workflow.",
            "Contact the administrator if the face image is unclear or attendance is rejected.",
        ],
    )
    add_heading(doc, "Device Operator Perspective", 1)
    add_bullets(
        doc,
        [
            "Power the ESP32-CAM with stable 5V.",
            "Ensure the laptop backend and ESP32-CAM are on the same network.",
            "Use the device status page to check current state and camera preview.",
            "For upload through mpremote, boot normally with IO0 disconnected from GND.",
            "For flashing through esptool, connect IO0 to GND and reset the board.",
        ],
    )
    doc.save(MONITORING_ROOT / "User Guidelines" / "User Guidelines.docx")


def installation_guidelines():
    doc = Document()
    add_heading(doc, "Installation Guidelines", 0)
    add_heading(doc, "Backend Setup", 1)
    add_bullets(
        doc,
        [
            "Open PowerShell in the attendance-system project folder.",
            "Create or use the Python 3.11 virtual environment: .venv311.",
            "Install dependencies from backend/requirements.txt.",
            "Use dlib-bin on Windows if dlib source compilation fails.",
            "Run the backend from the backend folder using python app.py.",
            "Open http://127.0.0.1:5000/login in a browser.",
        ],
    )
    doc.add_paragraph("Example commands:")
    doc.add_paragraph(
        "py -3.11 -m venv .venv311\n"
        ".\\.venv311\\Scripts\\python.exe -m pip install -r backend\\requirements.txt\n"
        "cd backend\n"
        "..\\.venv311\\Scripts\\python.exe app.py"
    )
    add_heading(doc, "ESP32-CAM Setup", 1)
    add_bullets(
        doc,
        [
            "Use ESP32-CAM camera firmware, not generic ESP32 firmware without camera support.",
            "For esptool flashing, connect IO0 to GND, reset, and keep IO0 grounded during flashing.",
            "For mpremote upload, remove IO0 from GND and reset into normal MicroPython boot.",
            "Upload camera-only files with firmware/upload_to_esp32.bat COM9.",
            "Connect ESP32-CAM to the same hotspot/router as the backend laptop.",
        ],
    )
    add_heading(doc, "Network Setup", 1)
    add_bullets(
        doc,
        [
            "When using laptop hotspot, backend address is normally http://192.168.137.1:5000.",
            "Set ESP32-CAM Wi-Fi SSID and password in firmware config or device_settings.json.",
            "Check the Devices page for heartbeat after reset.",
            "Open the ESP32-CAM IP address in a browser to see status and preview.",
        ],
    )
    doc.save(MONITORING_ROOT / "Installation Guidelines" / "Installation Guidelines.docx")


def testing_evaluation(evidence_lines, evidence_png):
    doc = Document()
    add_heading(doc, "Testing and Evaluation", 0)
    doc.add_paragraph("This document records the testing work performed for the Secure ESP32 Attendance System.")
    add_heading(doc, "Test Summary", 1)
    add_bullets(
        doc,
        [
            "Backend dependency import verification.",
            "Database availability and table count verification.",
            "Backend login route HTTP availability verification.",
            "ESP32-CAM configuration and upload workflow checks.",
            "Existing ESP32-CAM preview screenshots copied into evidence folder.",
        ],
    )
    add_heading(doc, "Automated Evidence", 1)
    for line in evidence_lines:
        doc.add_paragraph(line)
    if evidence_png.exists():
        doc.add_picture(str(evidence_png), width=Inches(6.5))

    add_heading(doc, "Manual Device Testing Checklist", 1)
    add_bullets(
        doc,
        [
            "Power ESP32-CAM with 5V and common GND.",
            "Confirm camera preview endpoint opens from browser.",
            "Activate attendance session and display QR code.",
            "Confirm ESP32-CAM state changes from WAIT_QR to VERIFY_QR.",
            "Look at camera and confirm face verification request reaches backend.",
            "Confirm attendance or verification event appears in backend pages.",
            "For separated fingerprint workflow, verify ESP32 Dev Board and R307S independently.",
        ],
    )

    for candidate in EVIDENCE_ROOT.glob("*.jpg"):
        try:
            doc.add_heading(candidate.name, level=2)
            doc.add_picture(str(candidate), width=Inches(5.5))
        except Exception:
            pass
    doc.save(MONITORING_ROOT / "Testing and Evaluation" / "Testing and Evaluation.docx")


def portfolio_index():
    doc = Document()
    add_heading(doc, "Portfolio Index", 0)
    doc.add_paragraph("Portfolio folder includes final presentation slides and supporting component files.")
    add_heading(doc, "Included Slides", 1)
    slides = sorted((PORTFOLIO_ROOT / "Final Presentation Slides").glob("*.pptx"))
    if slides:
        add_bullets(doc, [slide.name for slide in slides])
    else:
        doc.add_paragraph("No PPTX slides were found to copy.")
    add_heading(doc, "Demonstration Notes", 1)
    add_bullets(
        doc,
        [
            "Demonstrate administrator login and dashboard.",
            "Demonstrate student creation and face enrollment.",
            "Demonstrate session creation and QR generation.",
            "Demonstrate device heartbeat/status page.",
            "Demonstrate QR and face verification workflow.",
            "Explain separated ESP32-CAM and ESP32 Dev Board fingerprint architecture.",
        ],
    )
    doc.save(PORTFOLIO_ROOT / "Portfolio Index.docx")


def folder_readme():
    text = """Secure ESP32 Attendance System - Component Submission

Folder contents:
- Diagram screenshots and Requirements Catalog are in this Executing Folder.
- Implementation/Source Code contains the full project source and database files.
- Monitoring and Controlling contains User Guidelines, Installation Guidelines, and Testing and Evaluation.
- Portfolio contains final presentation slides and a portfolio index.

Zip format is used as required by the assignment.
"""
    (COMPONENT_ROOT / "README_Component_Submission.txt").write_text(text, encoding="utf-8")


def main():
    ensure_dirs()
    save_schema()
    evidence_lines, evidence_png = collect_test_evidence()
    requirements_catalog()
    user_guidelines()
    installation_guidelines()
    testing_evaluation(evidence_lines, evidence_png)
    portfolio_index()
    folder_readme()
    print(f"Component package documents generated at: {COMPONENT_ROOT}")


if __name__ == "__main__":
    main()
